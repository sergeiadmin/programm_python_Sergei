# подключение библиотеки 'tkinter'
from tkinter import *
from tkinter import messagebox

from tkinter import ttk
from tkinter import Label, Frame
from docx import Document # Для работы с документом
import subprocess # Для открытия сохраненного файла
from tkinter import Tk, Button

massiv_group = ("211", "411", "811", "1011", "1211")

Student_name_group = {
    "Анисимова Екатерина Денисовна" : "211",
    "Анненкова Александра Федоровна": "211",
    "Безбородых Варвара Сергеевна": "211",
    "Булаева Ксения Юрьевна": "211",
    "Васильченко Яна Александровна": "211",
    "Видюкова Алина Михайловна": "211",
    "Дегальцева Екатерина Сергеевна": "211",
    "Зюбан Анна Егоровна": "211",
    "Калабина Дарья Андреевна": "211",
    "Колосова Ангелина Ивановна": "211",
    "Кравченко Елизавета Николаевна": "211",
    "Левшина Лика Николаевна": "211",
    "Лемещенко Елизавета Сергеевна": "211",
    "Мироненко Елена Ивановна": "211",
    "Нарыкова Дарья Васильевна": "211",
    "Перебейнос Дарья Валерьевна": "211",
    "Пивовар Мария Анатольевна": "211",
    "Пчельникова София Андреевна": "211",
    "Смурыгина Татьяна Викторовна": "211",
    "Солодун Людмила Евгеньевна" : "211",
    "Тютюнникова Екатерина Сергеевна": "211",

    "Авраменко Мария Александровна": "411",
    "Афанасьева Виктория Игоревна": "411",
    "Бородина Виктория Федоровна": "411",
    "Бражникова Ксения Сергеевна": "411",
    "Веретенникова Дарья Ивановна": "411",
    "Гернеший Алина Викторовна": "411",
    "Гончарова Екатерина Владимировна": "411",
    "Дьякова Анна Максимовна": "411",
    "Копанева Вероника Романовна": "411",
    "Кравцова Мария Андреевна": "411",
    "Мышанская Алина Александровна": "411",
    "Санжарова Альбина Евгеньевна": "411",

    "Аношенков Тимур Алексеевич": "811",
    "Верещак Богдан Владимирович": "811",
    "Гезуля Денис Александрович": "811",
    "Гончаренко Владимир Владимирович": "811",
    "Коркодилов Никита Дмитриевич": "811",
    "Ливикин Даниил Вячеславович": "811",
    "Марков Иван Алексеевич": "811",
    "Пархомов Егор Сергеевич": "811",
    "Попов Трофим Иванович": "811",
    "Тоболенко Владислав Александрович": "811",
    "Трапезников Дмитрий Николаевич": "811",
    "Фомин Сергей Николаевич": "811",

    "Кузнецов Семен Игоревич": "1011",
    "Магомедов Гасан Ахмедович": "1011",
    "Малахов Матвей Алексеевмч": "1011",
    "Оксанченко Богдан Вячеславович": "1011",
    "Пересадченко Екатерина Дмитревна": "1011",
    "Уграватый Сергей Леонидович": "1011",

    "Аркатова Ирина Александровна": "1211",
    "Борисов Владислав Владимирович": "1211",
    "Бугаева Олеся Максимовна": "1211",
    "Жибоедов Денис Сергеевич": "1211",
    " Кальченко Артём Дмитриевич": "1211",
    " Лелюх Андрей Алексеевич": "1211"
}

#-------------------------------------------------------- form 1 ------------ Start block
# функция для расчета по нажатию кнопки
def run_batton():
    # метод 'get()' позволяет получить данные из полей
    a = str(tf.get())
    b = str(tf_2.get())
    # делаем условие
    if a == "user":
        if b == "123":
            messagebox.showinfo("Корректный ввод данных", "Пожалуйста, нажмите ОК для продолжения")
            newWindow_list()
            window.destroy()

        else:
            messagebox.showerror("Ошибка ввода данных", "Пожалуйста, введите корректные данные")
    else:
        messagebox.showerror("Ошибка ввода данных", "Пожалуйста, введите корректные данные")



window = Tk()

# название окна
window.title("Реестр заявок питания в столовой")
# размеры окна
window.geometry('1280x720')


# блок 'Frame'
frame = Frame(
    # параметр, указывающий на окно приложения
    window,

    # отступ по горизонтали в 15 пикселей
    padx=22,
    # отступ по вертикали в 15 пикселей
    pady=22

)
# позиционирование виджета в окне (expand - работа во всем окне)
frame.pack(expand=True)


# блок по добавлению надписи
lb = Label(
    # использование блока 'frame' в коде
    frame,
    text = "Логин: "
)
# метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
lb.grid(row=3, column=1)


# блок для введения пользовательской информации
tf = Entry(
    frame,
)
# указываем, что поле будет находиться справа от верхней надписи
tf.grid(row=3, column=2, ipadx=20, ipady=3, padx=5, pady=15)
# добавление второй надписи
lb_2 = Label(
    frame,
    text = "Пароль: "
)
lb_2.grid(row=5, column=1)

# добавление второго поля для ввода пользовательской информации
tf_2 = Entry(
    frame,
)
tf_2.grid(row=5, column=2, ipadx=20, ipady=3, padx=5, pady=15)

# добавление кнопки
btn = Button(
    frame,
    # подпись кнопки
    text="ОК",
    command=run_batton,
    bg = "#27AE60"
)
# ipadx=30, ipady=5 -- расширил кнопку
btn.grid(row=8, column=2,  ipadx=50, ipady=5)
#-------------------------------------------------------- form 1 ------------ End block

#-------------------------------------------------------- form 2 ------------ Start block

def newWindow_list():
    def exit_app():
       frame_Window_list.destroy()

    def avtor_app():
        window_avtor = Tk()
        window_avtor.title("Создатель приложения")
        window_avtor.geometry("350x50")

        # блок по добавлению надписи -------------------------------------- Заявка на питание:
        elToForm2_label_name_form = Label(
            # использование блока 'frame' в коде
            window_avtor,
            text="Разработчик программы: Любивый Сергей Николаевич"
        )
        # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
        elToForm2_label_name_form.grid(row=1, column=2)

    window2 = Tk()
    window2.title("Форма заполнения учета заявок на питание")
    window2.geometry("1280x720")
    # блок 'Frame'
    frame_Window_list = Frame(
        # параметр, указывающий на окно приложения
        window2,
        # отступ по горизонтали в 15 пикселей
        padx=22,
        # отступ по вертикали в 15 пикселей
        pady=22
    )

    # позиционирование виджета в окне (expand - работа во всем окне)
    frame_Window_list.pack(expand=True)

    main_menu = Menu(window2)
    window2.config(menu=main_menu)

    window2 = Menu( main_menu)
    main_menu.add_cascade(label="Форма заполнения учета заявок на питание", menu=window2)
    window2.add_command(label="Создатель программы", command=avtor_app)
    window2.add_command(label="Закрыть приложение", command=exit_app)

    # блок по добавлению надписи -------------------------------------- Форма заполнения учета заявок на питание
    elToForm2_label_name_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text = "Форма заполнения учета заявок на питание"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_name_form.grid(row=1, column=1)

    # блок по добавлению надписи -------------------------------------- Заявка на питание:
    elToForm2_label_name_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text="Заявка на питание:"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_name_form.grid(row=2, column=1)

    def add_info():
        group = combobox_group.get()
        student = combobox_student.get()
        info_text.insert(END, f"{group}, | {student}, - записан(a), ________ \n")

    def add_info_all():
        group = combobox_group.get()
        students = [student for student, student_group in Student_name_group.items() if student_group == group]
        for student in students:
            info_text.insert(END, f"{group}, | {student}, - записан(a), ________ \n")

    def delete_info():
        info_text.delete("end-2l", "end-1l")

    def delete_all_info():
        info_text.delete(1.0, END)

    delete_all_button = Button(frame_Window_list, text="Удалить все", command=delete_all_info, bg = "#9B5038")
    delete_all_button.grid(row=3, column=4, padx=1, pady=1, ipadx=20, ipady=3)

    add_all_button = Button(frame_Window_list, text="Добавить всех", command=add_info_all)
    add_all_button.grid(row=4, column=4, padx=1, pady=1, ipadx=20, ipady=3,)

    add_button = Button(frame_Window_list, text="Добавить запись", command=add_info, bg = "#27AE60")
    add_button.grid(row=4, column=5, padx=1, pady=1, ipadx=20, ipady=3,)

    delete_button = Button(frame_Window_list, text="Удалить посл. запись", command=delete_info)
    delete_button.grid(row=4, column=6, padx=1, pady=1, ipadx=20, ipady=3,)

    # Накопительно поле информации --------------------------------------
    info_text = Text(frame_Window_list, width=70, height=20)
    info_text.grid(row=3, column=1, padx=5, pady=5)



    # блок по добавлению надписи --------------------------------------
    elToForm2_label_data_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text = "Дата подачи заявки"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_data_form.grid(row=1, column=5)

    # Создаем поле ввода даты с маской
    date_entry = Entry(
        frame_Window_list
    )
    date_entry.grid(row=1, column=6, ipadx=20, ipady=3, padx=5, pady=15)

    # блок по добавлению надписи -------------------------------------- group_number Student's name
    elToForm2_label_group_number_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text="Номер группы:"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_group_number_form.grid(row=2, column=5)

    # добавление второго поля для ввода пользовательской информации --------------------------------------------------------- Список начало
    def on_select(event):
        selected_group = combobox_group.get()
        filtered_students = [student for student, group in Student_name_group.items() if group == selected_group]

        # Обновление значений второго выпадающего списка
        combobox_student['values'] = filtered_students

    # добавление второго поля для ввода пользовательской информации --------------------------------------------------------- Список начало
    combobox_group = ttk.Combobox(frame_Window_list, values=massiv_group)
    combobox_group.bind("<<ComboboxSelected>>", on_select)
    combobox_group.grid(row=2, column=6, ipadx=20, ipady=3, padx=5, pady=15)
    # добавление второго поля для ввода пользовательской информации --------------------------------------------------------- Список конец

    # блок по добавлению надписи -------------------------------------- Student's name
    elToForm2_label_Student_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text="ФИО Студента:"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_Student_form.grid(row=3, column=5)

    # добавление второго поля для ввода пользовательской информации ---------------------------------------------------------

    # добавление второго поля для ввода пользовательской информации --------------------------------------------------------- Список начало
    combobox_student = ttk.Combobox(frame_Window_list)
    combobox_student.grid(row=3, column=6, ipadx=20, ipady=3, padx=5, pady=15)
    # добавление второго поля для ввода пользовательской информации --------------------------------------------------------- Список конец

    # блок по добавлению надписи --------------------------------------  Разделительная линия:
    elToForm2_label_Line_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text="_________________________________________________"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_Line_form.grid(row=5, column=5)

    # блок по добавлению надписи --------------------------------------  Group curator:
    elToForm2_label_Group_curator_form = Label(
        # использование блока 'frame' в коде
        frame_Window_list,
        text="Фамилия (Инициалы) куратора:"
    )
    # метод 'grid' (текст будет в 3-ей строке, 1-ом столбце)
    elToForm2_label_Group_curator_form.grid(row=6, column=5)

    # добавление второго поля для ввода пользовательской информации ---------------------------------------------------------
    elToForm2_Edit_Group_curator_form = Entry(
        frame_Window_list,
    )
    elToForm2_Edit_Group_curator_form.grid(row=6, column=6, ipadx=20, ipady=3, padx=5, pady=15)

    #---------------------------------------------------------
    # добавление кнопки ))))) Подтвердить сведения ((((((
    #---------------------------------------------------------

    def create_word_document():
        if not info_text.get("1.0",
                             "end-1c") or not combobox_group.get() or not date_entry.get() or not elToForm2_Edit_Group_curator_form.get():
            messagebox.showerror("Ошибка ввода данных!", "Пожалуйста, заполните все поля для формирования заявки.")
        else:
            document = Document()
            document.add_heading('Список студентов на питание в столовую', 0)
            document.add_paragraph(f"Группа: \n {combobox_group.get()}")
            document.add_paragraph(f"ФИО студент(а): \n {info_text.get('1.0', 'end')}")
            document.add_paragraph(f"Куратор группы: {elToForm2_Edit_Group_curator_form.get()}                                     _________________ подпись\n ")
            document.add_paragraph(f"Дата: {date_entry.get()}")

            document.save(f"{date_entry.get()}_Заявка_на_питание.doc"
                          f"x")
            file_path = f"{date_entry.get()}_Заявка_на_питание.docx"
            subprocess.Popen(['start', file_path], shell=True)
            print("Документ создан успешно!")


    BtnToForm2_Btn_1 = Button(
        frame_Window_list,
        # подпись кнопки
        text="Подтвердить сведения",
        command=create_word_document,
        bg = "#27AE60"
    )
    # ipadx=30, ipady=5 -- расширил кнопку
    BtnToForm2_Btn_1.grid(row=8, column=6,  ipadx=80, ipady=5)

#-------------------------------------------------------- form 2 ------------ End block



#-------------------------------------------------------- Другое ------------ Start block

#-------------------------------------------------------- Другое ------------ Endblock
# запуск приложения
window.mainloop()